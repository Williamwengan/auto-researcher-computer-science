#!/usr/bin/env python3
from __future__ import annotations

import argparse
import re
from pathlib import Path


def add_paragraph_from_line(doc, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        doc.add_paragraph("")
        return
    if stripped.startswith("#"):
        level = min(len(stripped) - len(stripped.lstrip("#")), 4)
        text = stripped[level:].strip()
        doc.add_heading(text, level=level)
        return
    if stripped.startswith("- "):
        doc.add_paragraph(stripped[2:].strip(), style="List Bullet")
        return
    if re.match(r"^\d+\.\s+", stripped):
        doc.add_paragraph(re.sub(r"^\d+\.\s+", "", stripped), style="List Number")
        return
    doc.add_paragraph(stripped)


def convert(md_path: Path, docx_path: Path) -> None:
    from docx import Document

    doc = Document()
    lines = md_path.read_text(encoding="utf-8").splitlines()

    in_code = False
    table_buffer: list[str] = []

    def flush_table() -> None:
        nonlocal table_buffer
        if not table_buffer:
            return
        rows = []
        for raw in table_buffer:
            cells = [cell.strip(" `") for cell in raw.strip().strip("|").split("|")]
            if cells and not all(set(cell) <= {"-", ":"} for cell in cells):
                rows.append(cells)
        if rows:
            width = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=width)
            table.style = "Table Grid"
            for i, row in enumerate(rows):
                for j in range(width):
                    table.cell(i, j).text = row[j] if j < len(row) else ""
        table_buffer = []

    for line in lines:
        stripped = line.strip()
        if stripped.startswith("```"):
            flush_table()
            in_code = not in_code
            continue
        if in_code:
            doc.add_paragraph(line, style="No Spacing")
            continue
        if stripped.startswith("|") and stripped.endswith("|"):
            table_buffer.append(line)
            continue
        flush_table()
        add_paragraph_from_line(doc, line)

    flush_table()
    docx_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(docx_path)


def main() -> None:
    parser = argparse.ArgumentParser(description="Simple Markdown to DOCX converter for competition reports.")
    parser.add_argument("input_md", type=Path)
    parser.add_argument("output_docx", type=Path)
    args = parser.parse_args()
    convert(args.input_md, args.output_docx)
    print(f"Wrote {args.output_docx}")


if __name__ == "__main__":
    main()
