#!/usr/bin/env python3
from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path
from typing import Any


ROOT = Path(__file__).resolve().parents[2]

OUT_MD = ROOT / "competition_submission/COMPETITION_DELIVERABLES_INDEX_CN.md"
OUT_JSON = ROOT / "competition_submission/COMPETITION_DELIVERABLES_INDEX.json"
OUT_DOCX = ROOT / "competition_submission/COMPETITION_DELIVERABLES_INDEX_CN.docx"


def exists(path: str) -> str:
    return "exists" if (ROOT / path).exists() else "missing"


def item(stage: str, kind: str, path: str, role: str, priority: str = "core") -> dict[str, str]:
    return {
        "stage": stage,
        "kind": kind,
        "priority": priority,
        "path": path,
        "status": exists(path),
        "role": role,
    }


def build_items() -> list[dict[str, str]]:
    rows: list[dict[str, str]] = []

    # 0. Final entry points
    stage = "00 最终入口与总收束"
    rows += [
        item(stage, "report", "competition_submission/FINAL_WORKFLOW_END_TO_END_CLOSING_REPORT_CN.md", "最终主线报告：说明 workflow 已经从 idea generation 走到真实执行反馈。"),
        item(stage, "json", "competition_submission/FINAL_WORKFLOW_END_TO_END_CLOSING_REPORT.json", "最终主线报告结构化摘要。"),
        item(stage, "script", "focused_workflow/scripts/build_final_workflow_closing_report.py", "生成最终端到端收束报告。"),
        item(stage, "report", "competition_submission/FINAL_STORYLINE_FOR_COMPETITION_CN.md", "比赛叙事主线材料。", "support"),
        item(stage, "report", "competition_submission/AI_RESEARCH_WORKFLOW_FULL_PROGRESS_REPORT_CN.md", "完整阶段性进展报告。", "support"),
        item(stage, "docx", "competition_submission/AI_RESEARCH_WORKFLOW_FULL_PROGRESS_REPORT_CN.docx", "完整进展报告 DOCX 版本。", "support"),
    ]

    # 1. Task input and baseline-focused outputs
    stage = "01 任务输入与结构化 idea generation"
    rows += [
        item(stage, "script", "focused_workflow/scripts/run_focused_workflow_v0_2.sh", "运行 v0.2 focused workflow。", "support"),
        item(stage, "script", "focused_workflow/scripts/render_prompt.py", "根据 task_spec 渲染 prompt。", "support"),
        item(stage, "script", "focused_workflow/scripts/validate_outputs.py", "检查 baseline_cards/focused_ideas/experiment_plan 输出结构。", "support"),
        item(stage, "report", "competition_submission/V03_PIPELINE_ROBUSTNESS_REPORT_CN.md", "v0.3 pipeline 鲁棒性阶段报告。", "support"),
        item(stage, "report", "competition_submission/V04_IDEA_WORKFLOW_FINAL_EVALUATION_REPORT_CN.md", "v0.4 idea workflow 评价报告。", "support"),
        item(stage, "report", "competition_submission/IDEA_GENERATION_MODULE_CARD_CN.md", "idea generation 模块卡片，可作为答辩简介。"),
    ]

    # 2. Evidence retrieval and evidence-grounded ideation
    stage = "02 论文证据检索与 evidence-grounded ideation"
    rows += [
        item(stage, "script", "focused_workflow/scripts/retrieve_paper_evidence.py", "检索/整理 paper evidence。"),
        item(stage, "script", "focused_workflow/scripts/validate_paper_evidence.py", "检查 paper evidence 文件。", "support"),
        item(stage, "script", "focused_workflow/scripts/validate_evidence_grounding.py", "检查 idea 是否绑定 evidence。", "support"),
        item(stage, "script", "focused_workflow/scripts/repair_v05_evidence_grounded_ideas.py", "v0.5 evidence-grounded idea 修复。"),
        item(stage, "script", "focused_workflow/scripts/seed_indoor_scene_evidence.py", "为 Indoor3D 构造 seeded evidence bank；最终必须透明披露。"),
        item(stage, "report", "competition_submission/V05_PAPER_RETRIEVAL_EVIDENCE_BINDING_REPORT_CN.md", "paper retrieval 与 evidence binding 报告。"),
        item(stage, "report", "competition_submission/V05_EVIDENCE_GROUNDED_IDEA_RESULT_ANALYSIS_CN.md", "evidence-grounded idea 结果分析。"),
        item(stage, "report", "competition_submission/V05_TARGETED_REPAIR_BEFORE_AFTER_REPORT_CN.md", "targeted repair 前后对比报告。"),
    ]

    # 3. Repair and scoring
    stage = "03 idea repair、质量评分与机制一致性修复"
    rows += [
        item(stage, "script", "focused_workflow/scripts/evaluate_idea_quality.py", "自动 idea quality scoring。"),
        item(stage, "script", "focused_workflow/scripts/repair_low_quality_ideas.py", "低质量 idea 修复。", "support"),
        item(stage, "script", "focused_workflow/scripts/apply_physical_property_v2_repair.py", "物理属性 v2 二次修复核心脚本。"),
        item(stage, "script", "focused_workflow/scripts/repair_physical_v2_evidence_cards.py", "物理属性 v2 evidence cards 修复。"),
        item(stage, "script", "focused_workflow/scripts/repair_indoor3d_evidence_cards.py", "Indoor3D evidence cards 修复。"),
        item(stage, "report", "competition_submission/V07_PHYSICAL_V2_CLAIM_REPAIR_SUMMARY_CN.md", "物理属性 claim repair 总结。"),
        item(stage, "report", "competition_submission/V07_PHYSICAL_V2_EVIDENCE_CARD_REPAIR_SUMMARY_CN.md", "物理属性 evidence-card repair 总结。"),
        item(stage, "report", "competition_submission/V07_INDOOR3D_CLAIM_FORMAT_REPAIR_SUMMARY_CN.md", "Indoor3D claim format repair 总结。"),
        item(stage, "report", "competition_submission/V07_INDOOR3D_EVIDENCE_CARD_REPAIR_SUMMARY_CN.md", "Indoor3D evidence-card repair 总结。"),
    ]

    # 4. Multi-LLM judge and Si2024 benchmark
    stage = "04 多模型匿名盲评与 Si et al.-style benchmark"
    rows += [
        item(stage, "script", "focused_workflow/scripts/create_blind_ab_review_pack.py", "生成 blind A/B review pack。"),
        item(stage, "script", "focused_workflow/scripts/run_blind_ab_llm_reviewer.py", "运行 multi-LLM blind reviewer。"),
        item(stage, "script", "focused_workflow/scripts/summarize_blind_ab_reviews.py", "汇总 blind A/B 评审结果。"),
        item(stage, "script", "focused_workflow/scripts/run_si2024_blind_ab_reviewer.py", "运行 Si et al.-style blind reviewer。"),
        item(stage, "script", "focused_workflow/scripts/multi_llm_judge.py", "multi-LLM judge 通用脚本。", "support"),
        item(stage, "report", "competition_submission/V06_MULTI_LLM_BLIND_AB_EVALUATION_REPORT_V2_CN.md", "v0.6 multi-LLM blind A/B 评价报告。"),
        item(stage, "report", "competition_submission/SI2024_BENCHMARK_EVALUATION_REPORT_CN.md", "正式 Si et al.-style benchmark 报告。"),
        item(stage, "report", "outputs/si2024_three_task_benchmark_summary/SI2024_THREE_TASK_BENCHMARK_SUMMARY_CN.md", "三任务 Si-style benchmark 总表。"),
    ]

    # 5. Reference claim verification
    stage = "05 Reference claim verification"
    rows += [
        item(stage, "script", "focused_workflow/scripts/verify_reference_claims.py", "检查 claim 是否绑定真实 paper id 且被 title/abstract/card 支撑。"),
        item(stage, "report", "competition_submission/V07_REFERENCE_CLAIM_VERIFICATION_FINAL_SUMMARY_CN.md", "v0.7 reference claim verification 最终汇总。"),
        item(stage, "report", "competition_submission/V07_REFERENCE_CLAIM_VERIFICATION_SUMMARY_V2_CN.md", "v0.7 reference claim verification 中间汇总。", "support"),
    ]

    # 6. Core benchmark and final plan
    stage = "06 核心 benchmark 收束与最终研究方案"
    rows += [
        item(stage, "script", "focused_workflow/scripts/select_final_candidates_v0_8.py", "v0.8 benchmark robustness / candidate selector 历史脚本。", "support"),
        item(stage, "script", "focused_workflow/scripts/build_v09_idea_generation_core_benchmark_report.py", "生成 V09 idea generation core benchmark 报告。"),
        item(stage, "script", "focused_workflow/scripts/build_v10_final_research_plan_package.py", "生成 V10 final research plan package。"),
        item(stage, "report", "competition_submission/V09_IDEA_GENERATION_CORE_BENCHMARK_REPORT_CN.md", "核心 idea generation benchmark 报告。"),
        item(stage, "json", "competition_submission/V09_IDEA_GENERATION_CORE_BENCHMARK_SUMMARY.json", "V09 结构化摘要。"),
        item(stage, "report", "competition_submission/V10_FINAL_RESEARCH_PLAN_PACKAGE_CN.md", "最终研究方案包。"),
        item(stage, "json", "competition_submission/V10_FINAL_RESEARCH_PLAN_PACKAGE.json", "最终研究方案包结构化数据。"),
        item(stage, "schema", "competition_submission/FINAL_RESEARCH_PLAN_SCHEMA.json", "最终研究方案 schema。"),
    ]

    # 7. Execution planning
    stage = "07 实验执行规划与 IAD scaffold"
    rows += [
        item(stage, "script", "focused_workflow/scripts/build_v11_experiment_execution_plan.py", "生成 V11 experiment execution plan。"),
        item(stage, "script", "focused_workflow/scripts/build_v12_iad_mvp_scaffold.py", "生成 V12 IAD MVP 脚本骨架。"),
        item(stage, "report", "competition_submission/V11_EXPERIMENT_EXECUTION_PLAN_CN.md", "实验执行规划报告。"),
        item(stage, "json", "competition_submission/V11_EXPERIMENT_EXECUTION_PLAN.json", "实验执行规划结构化数据。"),
        item(stage, "schema", "competition_submission/EXPERIMENT_EXECUTION_PLAN_SCHEMA.json", "实验执行计划 schema。"),
        item(stage, "report", "competition_submission/V12_IAD_MVP_SCRIPT_SCAFFOLD_CN.md", "IAD MVP scaffold 说明。"),
        item(stage, "json", "competition_submission/V12_IAD_MVP_SCRIPT_SCAFFOLD.json", "IAD MVP scaffold 文件清单。"),
        item(stage, "script", "iad_mvp/scripts/check_env.py", "检查 IAD 环境和 MVTec 数据结构。"),
        item(stage, "script", "iad_mvp/scripts/prepare_mvtec_subset.py", "准备 MVTec split。"),
        item(stage, "script", "iad_mvp/scripts/prepare_iad_reference_manifest.py", "生成 IAD manifest。"),
        item(stage, "script", "iad_mvp/scripts/build_reference_bank.py", "构建 normal reference bank。"),
        item(stage, "script", "iad_mvp/scripts/run_iad_baselines.py", "运行 lightweight nearest-reference baseline。"),
        item(stage, "script", "iad_mvp/scripts/score_reference_consistency.py", "生成 reference consistency 决策。"),
        item(stage, "script", "iad_mvp/scripts/run_iad_negative_controls.py", "运行轻量负控制。"),
        item(stage, "script", "iad_mvp/scripts/evaluate_iad_agent.py", "生成 IAD execution metrics。"),
    ]

    # 8. IAD real-data execution feedback
    stage = "08 IAD 真实数据执行反馈与自动修复案例"
    rows += [
        item(stage, "script", "focused_workflow/scripts/build_v13_iad_smoke_test_report.py", "生成 V13 IAD data readiness + smoke test 报告。"),
        item(stage, "script", "focused_workflow/scripts/build_v14_iad_threshold_calibration.py", "生成 V14 阈值校准报告。"),
        item(stage, "script", "focused_workflow/scripts/build_v15_iad_multicategory_smoke_test_report.py", "生成 V15 三类别迁移 smoke test 报告。"),
        item(stage, "script", "focused_workflow/scripts/build_v16_iad_per_category_threshold_calibration.py", "生成 V16 类别感知阈值校准报告。"),
        item(stage, "script", "focused_workflow/scripts/build_v17_iad_category_constrained_retrieval_report.py", "生成 V17 类别约束检索/归一化报告。"),
        item(stage, "script", "focused_workflow/scripts/build_v18_iad_execution_feedback_repair_case.py", "生成 V18 execution-feedback repair case 总结。"),
        item(stage, "report", "competition_submission/V13_IAD_DATA_READINESS_AND_SMOKE_TEST_CN.md", "MVTec bottle 单类别 smoke test。"),
        item(stage, "report", "competition_submission/V14_IAD_THRESHOLD_CALIBRATION_CN.md", "bottle 阈值校准：accepted anomaly 0→51。"),
        item(stage, "report", "competition_submission/V15_IAD_MULTICATEGORY_SMOKE_TEST_CN.md", "三类别迁移发现全局阈值不鲁棒。"),
        item(stage, "report", "competition_submission/V16_IAD_PER_CATEGORY_THRESHOLD_CALIBRATION_CN.md", "类别感知阈值将 FPR 0.574257→0.009901。"),
        item(stage, "report", "competition_submission/V17_IAD_CATEGORY_CONSTRAINED_RETRIEVAL_CN.md", "类别约束检索与类别内归一化，小幅提升并暴露 feature 瓶颈。"),
        item(stage, "report", "competition_submission/V18_IAD_EXECUTION_FEEDBACK_REPAIR_CASE_CN.md", "IAD execution-feedback repair case 总结。"),
        item(stage, "json", "competition_submission/V18_IAD_EXECUTION_FEEDBACK_REPAIR_CASE.json", "V18 结构化摘要。"),
    ]

    # 9. IAD output tables and artifacts
    stage = "09 IAD 输出表与中间产物"
    rows += [
        item(stage, "data", "iad_mvp/data/mvtec_split.json", "bottle split。"),
        item(stage, "data", "iad_mvp/data/mvtec_split_3cat.json", "bottle/cable/capsule 三类别 split。"),
        item(stage, "data", "iad_mvp/data/iad_reference_manifest.jsonl", "bottle IAD manifest。"),
        item(stage, "data", "iad_mvp/data/iad_reference_manifest_3cat.jsonl", "三类别 IAD manifest。"),
        item(stage, "output", "iad_mvp/outputs/patchcore_baseline/iad_baseline_scores.csv", "bottle lightweight baseline scores。"),
        item(stage, "output", "iad_mvp/outputs/reference_consistency/iad_reference_consistency_scores.csv", "bottle reference consistency scores。"),
        item(stage, "table", "iad_mvp/outputs/tables/iad_agent_execution_metrics.csv", "bottle execution metrics。"),
        item(stage, "table", "iad_mvp/outputs/tables/iad_negative_control_report.csv", "bottle negative control report。"),
        item(stage, "table", "iad_mvp/outputs/tables/iad_threshold_sweep.csv", "V14 threshold sweep。"),
        item(stage, "table", "iad_mvp/outputs/tables/iad_threshold_recommended_decisions.csv", "V14 recommended decisions。"),
        item(stage, "output", "iad_mvp/outputs/patchcore_baseline_3cat/iad_baseline_scores.csv", "三类别 global baseline scores。"),
        item(stage, "output", "iad_mvp/outputs/reference_consistency_3cat/iad_reference_consistency_scores_calibrated.csv", "三类别 global threshold calibrated scores。"),
        item(stage, "table", "iad_mvp/outputs/tables_3cat/iad_agent_execution_metrics.csv", "三类别 global threshold metrics。"),
        item(stage, "table", "iad_mvp/outputs/tables_3cat/iad_negative_control_report_3cat.csv", "三类别 negative control report。"),
        item(stage, "table", "iad_mvp/outputs/tables_3cat/iad_per_category_threshold_recommendations.csv", "V16 per-category threshold recommendations。"),
        item(stage, "table", "iad_mvp/outputs/tables_3cat/iad_per_category_calibrated_metrics.csv", "V16 global vs per-category metrics。"),
        item(stage, "output", "iad_mvp/outputs/reference_consistency_3cat/iad_reference_consistency_scores_per_category_calibrated.csv", "V16 per-category calibrated decisions。"),
        item(stage, "table", "iad_mvp/outputs/tables_3cat/iad_v17_category_constrained_metrics.csv", "V17 category-constrained metrics。"),
        item(stage, "table", "iad_mvp/outputs/tables_3cat/iad_v17_category_constrained_threshold_recommendations.csv", "V17 category-constrained threshold recommendations。"),
        item(stage, "output", "iad_mvp/outputs/reference_consistency_3cat_category_constrained/iad_reference_consistency_scores.csv", "V17 category-constrained consistency scores。"),
    ]

    # 10. Raw benchmark outputs and held-out notes
    stage = "10 原始 benchmark 输出与 held-out 样本"
    rows += [
        item(stage, "output-dir", "outputs/si2024_three_task_benchmark_summary", "Si-style 三任务 benchmark 总目录。"),
        item(stage, "output-dir", "outputs/v06_blind_ab_review_iad_20260712_105111", "IAD blind A/B review 原始输出目录。", "support"),
        item(stage, "output-dir", "outputs/v06_blind_ab_review_physical_property_20260712_105111", "Physical Property blind A/B review 原始输出目录。", "support"),
        item(stage, "output-dir", "outputs/v06_blind_ab_review_physical_property_v2_20260712_163934", "Physical Property v2 blind A/B review 原始输出目录。", "support"),
        item(stage, "output-dir", "outputs/v06_blind_ab_review_indoor_scene_generation_20260712_130427", "Indoor3D blind A/B review 原始输出目录。", "support"),
        item(stage, "note", "focused_workflow/tasks/benchmark_cv/02_human_motion_generation.yaml", "02 Human Motion 当前作为 held-out sample，不作为完整闭环主证据。", "support"),
        item(stage, "note", "focused_workflow/tasks/benchmark_cv/04_3d_reconstruction.yaml", "04 3D Reconstruction 当前作为 held-out sample，不作为完整闭环主证据。", "support"),
    ]

    return rows


def group_by_stage(rows: list[dict[str, str]]) -> dict[str, list[dict[str, str]]]:
    grouped: dict[str, list[dict[str, str]]] = {}
    for row in rows:
        grouped.setdefault(row["stage"], []).append(row)
    return grouped


def md_table(rows: list[dict[str, str]], columns: list[str]) -> str:
    header = "| " + " | ".join(columns) + " |"
    sep = "| " + " | ".join(["---"] * len(columns)) + " |"
    body = ["| " + " | ".join(str(row.get(col, "")) for col in columns) + " |" for row in rows]
    return "\n".join([header, sep, *body])


def build_markdown(payload: dict[str, Any]) -> str:
    parts = [
        "# 比赛交付物目录（按 Workflow 环节整理）",
        "",
        f"生成时间：{payload['generated_at']}",
        "",
        "生成脚本：`focused_workflow/scripts/build_competition_deliverables_index.py`",
        "",
        "## 使用说明",
        "",
        "这份目录不是新的实验报告，而是给比赛提交/答辩使用的索引。它把报告、脚本、输出表和中间产物按 workflow 环节整理，避免材料多而散。",
        "",
        "优先阅读顺序：",
        "",
        "1. `FINAL_WORKFLOW_END_TO_END_CLOSING_REPORT_CN.md`",
        "2. `V09_IDEA_GENERATION_CORE_BENCHMARK_REPORT_CN.md`",
        "3. `V10_FINAL_RESEARCH_PLAN_PACKAGE_CN.md`",
        "4. `V18_IAD_EXECUTION_FEEDBACK_REPAIR_CASE_CN.md`",
        "5. `SI2024_BENCHMARK_EVALUATION_REPORT_CN.md` 和 `V07_REFERENCE_CLAIM_VERIFICATION_FINAL_SUMMARY_CN.md`",
        "",
        "## 总览统计",
        "",
        md_table(payload["summary_rows"], ["metric", "value"]),
        "",
        "## 分环节交付物目录",
        "",
    ]

    grouped = group_by_stage(payload["items"])
    for stage, rows in grouped.items():
        parts += [
            f"### {stage}",
            "",
            md_table(rows, ["kind", "priority", "status", "path", "role"]),
            "",
        ]

    parts += [
        "## 边界提醒",
        "",
        "- 这份目录替代早期偏 Human Motion demo 的旧 tracker；旧 tracker 可作为历史记录，不建议作为当前主目录。",
        "- 当前主线是跨任务 AI 科研自动化 workflow，不是单个 IAD/CV 算法。",
        "- IAD V1.3–V1.8 是真实执行反馈案例；当前 IAD 结果仍是 lightweight scaffold，不是 PatchCore/anomalib 正式 benchmark。",
        "- Indoor3D 使用 seeded evidence bank，最终材料必须透明披露。",
        "- 02 Human Motion 和 04 3D Reconstruction 当前作为 held-out samples，不作为完整闭环主证据。",
        "",
    ]
    return "\n".join(parts)


def build_docx(payload: dict[str, Any], md_text: str) -> None:
    from docx import Document
    from docx.shared import Inches

    doc = Document()
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.55)
        section.right_margin = Inches(0.55)
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)

    doc.add_heading("比赛交付物目录（按 Workflow 环节整理）", 0)
    doc.add_paragraph(f"生成时间：{payload['generated_at']}")
    doc.add_paragraph("用途：比赛提交/答辩材料索引；按 workflow 环节整理报告、脚本、输出表和中间产物。")

    doc.add_heading("优先阅读顺序", level=1)
    for text in [
        "FINAL_WORKFLOW_END_TO_END_CLOSING_REPORT_CN.md",
        "V09_IDEA_GENERATION_CORE_BENCHMARK_REPORT_CN.md",
        "V10_FINAL_RESEARCH_PLAN_PACKAGE_CN.md",
        "V18_IAD_EXECUTION_FEEDBACK_REPAIR_CASE_CN.md",
        "SI2024_BENCHMARK_EVALUATION_REPORT_CN.md / V07_REFERENCE_CLAIM_VERIFICATION_FINAL_SUMMARY_CN.md",
    ]:
        doc.add_paragraph(text, style="List Number")

    doc.add_heading("总览统计", level=1)
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "metric"
    hdr[1].text = "value"
    for row in payload["summary_rows"]:
        cells = table.add_row().cells
        cells[0].text = row["metric"]
        cells[1].text = str(row["value"])

    grouped = group_by_stage(payload["items"])
    doc.add_heading("分环节交付物目录", level=1)
    for stage, rows in grouped.items():
        doc.add_heading(stage, level=2)
        table = doc.add_table(rows=1, cols=5)
        hdr = table.rows[0].cells
        for idx, col in enumerate(["kind", "priority", "status", "path", "role"]):
            hdr[idx].text = col
        for row in rows:
            cells = table.add_row().cells
            cells[0].text = row["kind"]
            cells[1].text = row["priority"]
            cells[2].text = row["status"]
            cells[3].text = row["path"]
            cells[4].text = row["role"]

    doc.add_heading("边界提醒", level=1)
    for text in [
        "这份目录替代早期偏 Human Motion demo 的旧 tracker；旧 tracker 可作为历史记录，不建议作为当前主目录。",
        "当前主线是跨任务 AI 科研自动化 workflow，不是单个 IAD/CV 算法。",
        "IAD V1.3–V1.8 是真实执行反馈案例；当前 IAD 结果仍是 lightweight scaffold，不是 PatchCore/anomalib 正式 benchmark。",
        "Indoor3D 使用 seeded evidence bank，最终材料必须透明披露。",
        "02 Human Motion 和 04 3D Reconstruction 当前作为 held-out samples，不作为完整闭环主证据。",
    ]:
        doc.add_paragraph(text, style="List Bullet")

    doc.save(OUT_DOCX)


def main() -> None:
    items = build_items()
    summary_rows = [
        {"metric": "indexed_items", "value": len(items)},
        {"metric": "core_items", "value": sum(1 for row in items if row["priority"] == "core")},
        {"metric": "support_items", "value": sum(1 for row in items if row["priority"] == "support")},
        {"metric": "missing_items", "value": sum(1 for row in items if row["status"] == "missing")},
        {"metric": "docx_available", "value": "yes"},
    ]
    payload = {
        "version": "competition_deliverables_index",
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "items": items,
        "summary_rows": summary_rows,
        "outputs": {
            "markdown": str(OUT_MD.relative_to(ROOT)),
            "json": str(OUT_JSON.relative_to(ROOT)),
            "docx": str(OUT_DOCX.relative_to(ROOT)),
        },
        "boundary": "Index only; no new experimental claim.",
    }

    OUT_JSON.write_text(json.dumps(payload, ensure_ascii=False, indent=2), encoding="utf-8")
    md_text = build_markdown(payload)
    OUT_MD.write_text(md_text, encoding="utf-8")
    build_docx(payload, md_text)

    print(f"Wrote {OUT_MD}")
    print(f"Wrote {OUT_JSON}")
    print(f"Wrote {OUT_DOCX}")
    print(
        "Summary: "
        f"indexed_items={summary_rows[0]['value']}, "
        f"missing_items={summary_rows[3]['value']}"
    )


if __name__ == "__main__":
    main()
